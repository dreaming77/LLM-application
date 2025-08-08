import os
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
import re
from docx import Document
from typing import List, Dict, Tuple, Optional
import pdfplumber

# 1. 文本加载
def load_docx_text(file_path: str) -> str:
    """
    加载DOCX文件内容为纯文本。
    提取docx文档的主要工作有：对标题的加粗格式保留和处理表格内容。

    参数:
        file_path: DOCX文件路径

    返回:
        文档纯文本内容
    """
    doc = Document(file_path)
    full_text = []

    for para in doc.paragraphs:
        # 保留加粗标题格式
        if any(run.bold for run in para.runs):
            full_text.append(f"**{para.text.strip()}**")
        else:
            full_text.append(para.text.strip())

    # 处理表格内容（刑法文档中的表格较少，但需考虑）
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                row_text.append(cell.text.strip())
            full_text.append(" | ".join(row_text))

    print(f"\nDOCX文本提取完成")
    return "\n".join(full_text)


def load_pdf_text(pdf_path: str) -> str:
    """
    从PDF文件中提取文本并添加页码标记

    参数:
        pdf_path: PDF文件路径

    返回:
        带有页码标记的完整文本字符串
    """
    full_text = ""
    with pdfplumber.open(pdf_path) as pdf:
        total_pages = len(pdf.pages)

        for i, page in enumerate(pdf.pages):
            # 提取当前页文本
            page_text = page.extract_text()

            # 添加页码标记 (格式: ===== Page X =====)
            page_marker = f"\n\n{'=' * 5} Page {i + 1} {'=' * 5}\n\n"

            # 第一页前不添加空行
            if i == 0:
                full_text += page_marker.strip() + "\n"
            else:
                full_text += page_marker

            full_text += page_text + "\n"

            # 打印进度
            if (i + 1) % 10 == 0 or (i + 1) == total_pages:
                print(f"已处理 {i + 1}/{total_pages} 页")

    print(f"\nPDF文本提取完成，共 {total_pages} 页")
    return full_text


# 2. 文本分割
def dynamic_law_docx_splitter(text: str, max_chunk_size: int = 1000) -> List[str]:
    """
    基于语义边界动态切分法律文本，保留编/章/节/条/款的结构完整性

    参数:
        text: 完整的法律文本
        max_chunk_size: 最大块大小（字符数）

    返回:
        切分后的文本块列表
    """
    # 结构层级正则表达式
    patterns = [
        (r"\*\*第[零一二三四五六七八九十百]+编\s*[^*]+\*\*", "编"),  # 编
        (r"\*\*第[零一二三四五六七八九十百]+章\s*[^*]+\*\*", "章"),  # 章
        (r"\*\*第[零一二三四五六七八九十百]+节\s*[^*]+\*\*", "节"),  # 节
        (r"第[零一二三四五六七八九十百]+条\s*[^】]+", "条")  # 条
    ]

    chunks = []
    current_chunk = ""
    last_delimiter = ""

    # 按行处理保持结构
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # 检查结构标记
        found_delimiter = None
        for pattern, delimiter in patterns:
            if re.search(pattern, line):
                found_delimiter = delimiter
                break

        # 遇到新结构时处理
        if found_delimiter:
            # 保存当前块
            if current_chunk:
                chunks.append(current_chunk.strip())
                current_chunk = ""

            # 根据层级添加分隔符
            if found_delimiter in ("编", "章", "节"):
                chunks.append(line)
                last_delimiter = found_delimiter
                continue
            elif found_delimiter == "条":
                current_chunk = line
                last_delimiter = found_delimiter
                continue

        # 普通内容行
        current_chunk += "\n" + line

        # 长度检查（排除章节标题）
        if len(current_chunk) > max_chunk_size and last_delimiter == "条":
            chunks.append(current_chunk.strip())
            current_chunk = ""

    # 添加最后一块
    if current_chunk:
        chunks.append(current_chunk.strip())

    return chunks


def dynamic_law_pdf_splitter(text: str, max_length: int = 1000) -> List[Dict[str, str]]:
    """
    基于语义边界动态切分民法典文本

    参数:
        text: 完整的民法典文本
        max_length: 单个片段的最大长度限制

    返回:
        包含元数据的文本片段列表
    """
    # 预处理：按页码分割文本
    page_splits = re.split(r'={5,}\s*Page\s+(\d+)\s*={5,}', text)
    pages = []
    for i in range(1, len(page_splits), 2):
        pages.append({
            "page_num": int(page_splits[i]),
            "content": page_splits[i + 1].strip()
        })

    # 结构解析状态变量
    current_chapter = None
    current_section = None
    current_article = None
    segments = []

    # 辅助函数：创建新片段
    def create_segment(content: str,
                       page_num: int,
                       chapter: Optional[str],
                       section: Optional[str],
                       article: Optional[str]) -> Dict[str, str]:
        """创建带元数据的文本片段"""
        return {
            "text": content.strip(),
            "metadata": {
                "page": page_num,
                "chapter": chapter,
                "section": section,
                "article": article,
                "segment_type": "article" if article else "section" if section else "chapter"
            }
        }

    # 解析正则表达式
    chapter_pattern = re.compile(r'^第([一二三四五六七八九十百]+)章\s+(.+)$')
    section_pattern = re.compile(r'^第([一二三四五六七八九十百]+)节\s+(.+)$')
    article_pattern = re.compile(r'^第([一二三四五六七八九十百]+)条\s*(.+)?$')
    page_header_pattern = re.compile(r'^={5,}\s*Page\s+\d+\s*={5,}$')

    # 按页处理文本
    for page in pages:
        lines = page["content"].split('\n')
        buffer = []  # 当前片段缓冲区
        buffer_page = page["page_num"]
        buffer_chapter = current_chapter
        buffer_section = current_section
        buffer_article = current_article

        for line in lines:
            line = line.strip()
            if not line or page_header_pattern.match(line):
                continue  # 跳过空行和页眉

            # 检查章节变化
            chapter_match = chapter_pattern.match(line)
            section_match = section_pattern.match(line)
            article_match = article_pattern.match(line)

            # 处理章节标题
            if chapter_match:
                # 保存上一个片段
                if buffer:
                    segments.append(create_segment(
                        "\n".join(buffer),
                        buffer_page,
                        buffer_chapter,
                        buffer_section,
                        buffer_article
                    ))
                    buffer = []

                # 更新章节信息
                current_chapter = f"第{chapter_match.group(1)}章 {chapter_match.group(2)}"
                current_section = None
                current_article = None
                buffer_chapter = current_chapter
                buffer_section = None
                buffer_article = None
                buffer_page = page["page_num"]
                buffer.append(line)
                continue

            # 处理节标题
            if section_match:
                # 保存上一个片段
                if buffer:
                    segments.append(create_segment(
                        "\n".join(buffer),
                        buffer_page,
                        buffer_chapter,
                        buffer_section,
                        buffer_article
                    ))
                    buffer = []

                # 更新节信息
                current_section = f"第{section_match.group(1)}节 {section_match.group(2)}"
                current_article = None
                buffer_section = current_section
                buffer_article = None
                buffer_page = page["page_num"]
                buffer.append(line)
                continue

            # 处理条文
            if article_match:
                # 保存上一个条文片段
                if buffer and buffer_article:
                    segments.append(create_segment(
                        "\n".join(buffer),
                        buffer_page,
                        buffer_chapter,
                        buffer_section,
                        buffer_article
                    ))
                    buffer = []

                # 开始新条文
                article_num = f"第{article_match.group(1)}条"
                current_article = article_num
                buffer_article = article_num
                buffer_page = page["page_num"]

                # 如果条文有内容，直接添加
                if article_match.group(2):
                    buffer.append(f"{article_num} {article_match.group(2)}")
                else:
                    buffer.append(article_num)
                continue

            # 普通内容行
            buffer.append(line)

            # 检查缓冲区长度
            current_length = sum(len(l) for l in buffer)
            if current_length > max_length and buffer_article:
                # 对长条文进行适度分割
                segments.append(create_segment(
                    "\n".join(buffer),
                    buffer_page,
                    buffer_chapter,
                    buffer_section,
                    buffer_article
                ))
                buffer = []
                buffer_article = f"{buffer_article}(续)"

        # 处理页面末尾的缓冲区
        if buffer:
            segments.append(create_segment(
                "\n".join(buffer),
                buffer_page,
                buffer_chapter,
                buffer_section,
                buffer_article
            ))

    return segments


# 3. 创建嵌入模型 - 使用本地模型
def create_embeddings():
    # 修改为你的本地模型路径
    model_path = "./text2vec-large-chinese"  # 模型下载在此目录
    
    # 检查模型是否存在
    if not os.path.exists(model_path):
        raise FileNotFoundError(f"找不到嵌入模型目录: {model_path}")
    
    embeddings = HuggingFaceEmbeddings(
        model_name=model_path,  # 使用本地路径
        model_kwargs={'device': 'cpu'},  # 使用GPU可改为 'cuda'
        encode_kwargs={'normalize_embeddings': False}
    )
    print(f"本地嵌入模型加载成功: {model_path}")
    return embeddings


# 4. 创建向量存储
def create_vector_db(split_docs, embeddings):

    # 初始化一个空的FAISS索引（先创建一个虚拟文本）
    empty_text = "初始化"
    vector_db = FAISS.from_texts(
        texts=[empty_text],
        embedding=embeddings,
        metadatas=[{}]  # 空元数据
    )

    # 优化索引参数
    index = vector_db.index
    index.nprobe = 10  # 平衡速度与精度

    # 转换为字典列表（添加元数据）
    docx_records = [
        {"text": chunk, "source": "docx", "page": None}
        for chunk in split_docs[0]
    ]
    # 添加到FAISS
    vector_db.add_texts(
        texts=[r["text"] for r in docx_records],
        metadatas=[{"source": r["source"], "page": r["page"]} for r in docx_records]
    )

    # 确保PDF的字典中有"text"键
    for chunk in split_docs[1]:
        if "content" in chunk:
            chunk["text"] = chunk.pop("content")

    # 添加到FAISS
    vector_db.add_texts(
        texts=[r["text"] for r in split_docs[1]],
        metadatas=[{"source": "pdf", "page": r.get("page")} for r in split_docs[1]]
    )

    # 保存向量库到本地
    save_path = "faiss_vector_db"
    vector_db.save_local(save_path)
    print(f"向量库已保存至 {save_path}，包含 {vector_db.index.ntotal} 个向量")

    return vector_db


if __name__ == "__main__":
    docx_path = './dataset/刑法.docx'
    pdf_path = './dataset/民法典.pdf'
    # 步骤执行
    print("=" * 50 + "\n法律文档向量化处理开始\n" + "=" * 50)

    # 1. 加载文档
    raw_docs_docx = load_docx_text(docx_path)
    raw_docs_pdf = load_pdf_text(pdf_path)

    # 2. 分割文本
    split_docs_docx = dynamic_law_docx_splitter(raw_docs_docx)
    split_docs_pdf = dynamic_law_pdf_splitter(raw_docs_pdf)
    split_docs = [split_docs_docx, split_docs_pdf]

    # 3. 初始化嵌入模型
    embeddings = create_embeddings()

    # 4. 创建并保存向量库
    vector_db = create_vector_db(split_docs, embeddings)

