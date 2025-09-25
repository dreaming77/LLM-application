import re
from typing import List
from langchain.schema import Document
from langchain.text_splitter import TextSplitter
from PyPDF2 import PdfReader
import docx


class LawTextSplitter(TextSplitter):
    """自定义法律文本分割器，按章/节/条结构分割"""

    def __init__(self, law_name: str):
        super().__init__()
        self.law_name = law_name

    def split_text(self, text: str) -> List[Document]:
        # 预处理文本：合并换行和多余空格
        text = re.sub(r'\s+', ' ', text)

        # 按章分割
        chapters = re.split(r'(第[一二三四五六七八九十百千]+章\s*[^第条节]*)', text)
        chapters = [ch.strip() for ch in chapters if ch.strip()]

        docs = []
        current_chapter = ""
        current_section = ""

        for i, chunk in enumerate(chapters):
            # 检测章
            if re.match(r'第[一二三四五六七八九十百千]+章', chunk):
                # 只提取章标题，不包含后续内容
                chapter_match = re.match(r'(第[一二三四五六七八九十百千]+章\s*[^第条节]*)', chunk)
                if chapter_match:
                    current_chapter = chapter_match.group(1).strip()
                else:
                    current_chapter = chunk
                current_section = ""  # 重置节
                continue

            # 检测节（如果有）
            section_match = re.search(r'(第[一二三四五六七八九十百千]+节\s*[^第条]*)', chunk)
            if section_match:
                current_section = section_match.group(1).strip()
                # 从chunk中移除节标题
                chunk = re.sub(r'第[一二三四五六七八九十百千]+节\s*[^第条]*', '', chunk)

            # 在当前章/节中按条分割
            articles = re.split(r'(第[一二三四五六七八九十百千]+条)', chunk)
            articles = [art.strip() for art in articles if art.strip()]

            current_article_id = ""
            for art in articles:
                if re.match(r'第[一二三四五六七八九十百千]+条', art):
                    current_article_id = art
                    continue

                if current_article_id and art:
                    # 创建文档块
                    doc = Document(
                        page_content=art,
                        metadata={
                            "law_name": self.law_name,
                            "chapter_title": current_chapter,
                            "section_title": current_section,
                            "article_id": current_article_id
                        }
                    )
                    docs.append(doc)

        return docs


def read_pdf(file_path):
    """读取PDF文件"""
    reader = PdfReader(file_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text


def read_docx(file_path):
    """读取DOCX文件"""
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])


def read_txt(file_path):
    """读取TXT文件"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()


def process_documents():
    """处理所有法律文档并创建向量存储"""
    from config.settings import RAW_DATA_DIR, VECTOR_STORE_PATH, EMBEDDING_MODEL_PATH
    from langchain_huggingface import HuggingFaceEmbeddings
    from langchain_community.vectorstores import FAISS
    import os

    # 读取文件
    civil_code_text = read_pdf(os.path.join(RAW_DATA_DIR, "civil_code.pdf"))
    penal_code_text = read_txt(os.path.join(RAW_DATA_DIR, "penal_code.txt"))
    corporation_law_text = read_docx(os.path.join(RAW_DATA_DIR, "corporation_law.docx"))

    civil_code_text = read_pdf(civil_code_text)
    penal_code_text = read_txt(penal_code_text)
    corporation_law_text = read_docx(corporation_law_text)

    # 分割文本
    splitter_civil = LawTextSplitter(law_name="民法典")
    splitter_penal = LawTextSplitter(law_name="刑法")
    splitter_corp = LawTextSplitter(law_name="公司法")

    docs_civil = splitter_civil.split_text(civil_code_text)
    docs_penal = splitter_penal.split_text(penal_code_text)
    docs_corp = splitter_corp.split_text(corporation_law_text)

    all_docs = docs_civil + docs_penal + docs_corp
    print(f"总共分割出 {len(all_docs)} 个条文")

    # 创建向量存储
    embeddings = HuggingFaceEmbeddings(
        model_name=EMBEDDING_MODEL_PATH,
        model_kwargs={'device': 'cuda:5'}  # 使用GPU可改为 'cuda'
    )
    vectorstore = FAISS.from_documents(all_docs, embeddings)

    # 保存索引
    vectorstore.save_local(VECTOR_STORE_PATH)
    print(f"向量存储已保存到 {VECTOR_STORE_PATH}")

    return vectorstore
